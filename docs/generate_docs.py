import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Set page margins
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Title
title = doc.add_heading('ElewaSTEM (Mwalimu STEM) — Hackathon Demo Video Guide', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p_sub = doc.add_paragraph('Official 3-Minute Presentation Script & Screen Setup for All Things Agentic Hackathon')
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.runs[0].font.italic = True
p_sub.runs[0].font.color.rgb = RGBColor(16, 120, 80)

doc.add_paragraph()

# Section 1: Screen Setup
doc.add_heading('1. Screen & Browser Setup Checklist', level=1)
doc.add_paragraph('1. Open your browser to: http://localhost:8000')
doc.add_paragraph('2. Zoom in to 110% or 125% (Ctrl + +) so all text, badges, and diagrams are large and crisp.')
doc.add_paragraph('3. Hide the bookmarks bar (Ctrl + Shift + B) and close unrelated tabs.')
doc.add_paragraph('4. Header settings: Language = Kiswahili (or Sheng), Region = Bonde la Ziwa (Kisumu), Mode = Hadithi (Temp 0.75).')

doc.add_paragraph()

# Section 2: Timeline Table
doc.add_heading('2. The 3-Minute Video Timeline', level=1)

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Scene / Time'
hdr_cells[1].text = 'What to Show on Screen'
hdr_cells[2].text = 'Core Focus'

data = [
    ('Scene 1 (0:00 - 0:35)', 'App Homepage & Welcoming Banner', 'The Problem: 89% Learning Poverty & 80% Language Gap'),
    ('Scene 2 (0:35 - 1:25)', 'Chat, Voice, SVG Diagram, Tactile/Sign', 'Learner Experience: Socratic friend, Vector SVG, Inclusion'),
    ('Scene 3 (1:25 - 2:05)', 'Stakeholder Hub (Wazazi & Walimu)', '360° Network: Remote Parent 2G SMS & CBC Lesson Plans'),
    ('Scene 4 (2:05 - 2:40)', 'DPA Legal & Ethics Hub', 'Google Cloud Run, 8+ DPAs, ETHOS/OASIS, Maasai Swarm'),
    ('Scene 5 (2:40 - 3:00)', 'Main Chat / GitHub Repo', 'Outro & Call to Action: "Usimeze, Elewa!"')
]

for time_col, screen_col, focus_col in data:
    row_cells = table.add_row().cells
    row_cells[0].text = time_col
    row_cells[1].text = screen_col
    row_cells[2].text = focus_col

doc.add_paragraph()

# Section 3: Word-for-Word Script
doc.add_heading('3. Word-for-Word Narration Script', level=1)

doc.add_heading('Scene 1: The Hook & The Problem (0:00 – 0:35)', level=2)
doc.add_paragraph('"Hello judges! According to the World Bank and UNESCO, 89% of 10-year-old children in Sub-Saharan Africa cannot read and understand a simple text. Furthermore, 80% of African children are taught in a foreign language they do not speak at home, and over 60% of rural schools have zero internet connectivity.\n\nMeet ElewaSTEM (Mwalimu STEM) — an offline-first, voice-enabled AI learning ecosystem powered by Google Cloud and Gemini that bridges language, disability, and infrastructure barriers across 16+ African languages."')

doc.add_heading('Scene 2: Core Learner Experience & SVG Diagrams (0:35 – 1:25)', level=2)
doc.add_paragraph('"Instead of generic Western examples, ElewaSTEM acts as a loving, encouraging friend. When explaining photosynthesis in Kisumu, it uses Lake Victoria water hyacinth and indigenous Osuga greens.\n\nFor visual learners, it generates 100% offline, zero-bandwidth vector SVG diagrams. For blind students, it provides sensory Tactile Audio descriptions, and for deaf learners, structured Sign Language cues — ensuring true multi-sensory inclusion."')

doc.add_heading('Scene 3: 360° Stakeholder Network & Remote Parent Sync (1:25 – 2:05)', level=2)
doc.add_paragraph('"Education is a community effort. ElewaSTEM solves the parent-teacher gap through our 360° Stakeholder Loop.\n\nParents working away from home receive automated 2G feature phone SMS progress updates without needing a smartphone. Teachers receive instant CBC curriculum-aligned lesson plans, and community mentors get zero-cost STEM club guides."')

doc.add_heading('Scene 4: Google Cloud Architecture & Data Sovereignty (2:05 – 2:40)', level=2)
doc.add_paragraph('"Under the hood, ElewaSTEM runs on Google Cloud Run, Google Cloud Build, and Gemini 2.5 Flash on Vertex AI.\n\nWe enforce strict digital sovereignty across 8 African Data Protection Acts and the AU Malabo Convention. With our OASIS protocol, all child data is processed 100% on-device on the edge with zero persistent cloud tracking, guarded by a Maasai-elder calibrated autonomous multi-agent swarm."')

doc.add_heading('Scene 5: Outro & Call to Action (2:40 – 3:00)', level=2)
doc.add_paragraph('"ElewaSTEM proves that frontier AI can celebrate African languages, protect children\'s data rights, and bring world-class STEM education to the most remote villages.\n\nOur motto is: \'Usimeze, Elewa!\' — Don\'t cram, understand! Thank you!"')

doc.add_paragraph()

# Section 4: Links
doc.add_heading('4. Important Project Links', level=1)
doc.add_paragraph('• GitHub Repository: https://github.com/ondinookello96/elewastem')
doc.add_paragraph('• Devpost Project: https://devpost.com/software/elewastem-mwalimu-stem')
doc.add_paragraph('• Local Live App: http://localhost:8000')

out_path = r'C:\Users\Cosmas\.gemini\antigravity\scratch\elewastem\docs\ElewaSTEM_Demo_Video_Script_and_Guide.docx'
doc.save(out_path)
print(f'Successfully saved to {out_path}')
